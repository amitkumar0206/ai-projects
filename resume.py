import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Function to add styles
def set_style(paragraph, size, bold=False):
    run = paragraph.add_run()
    run.font.size = Pt(size)
    run.bold = bold

def convert_html_to_word(html_file, word_file):
    # Read the HTML file
    with open(html_file, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f, 'html.parser')

    # Create a new Document
    doc = Document()

    # Title
    title = soup.title.string if soup.title else "Document"
    doc.add_heading(title, 0)

    # Process body content
    for element in soup.body.find_all(True):  # Find all tags
        if element.name == 'div' and 'header' in element.get('class', []):
            # Header section
            if element.h1:
                doc.add_paragraph(element.h1.text, style='Title')
            else:
                doc.add_paragraph("No header title found", style='Title')
            contact_info = element.find(class_='contact-info')
            if contact_info:
                doc.add_paragraph(contact_info.text)
            doc.add_paragraph().add_run()  # Blank line

        elif element.name in ['h1', 'h2', 'h3', 'h4']:
            doc.add_heading(element.text.strip(), level=int(element.name[1]))

        elif element.name == 'p':
            doc.add_paragraph(element.text.strip())

        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(f'- {li.text.strip()}')

        elif element.name == 'table':
            table = doc.add_table(rows=1, cols=len(element.find_all('th')))
            hdr_cells = table.rows[0].cells
            for idx, th in enumerate(element.find_all('th')):
                hdr_cells[idx].text = th.text.strip()
            for tr in element.find_all('tr'):
                row_cells = table.add_row().cells
                for idx, td in enumerate(tr.find_all('td')):
                    row_cells[idx].text = td.text.strip()

    # Save the Word document
    doc.save(word_file)

# Example usage
html_file = 'Resume.html'
word_file = 'Resume.docx'
convert_html_to_word(html_file, word_file)